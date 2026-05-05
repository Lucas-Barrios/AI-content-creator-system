"""
app.py — SRH University AI Content Creator
Full-viewport demo UI · Run: streamlit run app.py
"""

import io
import os
import re
import streamlit as st
from dotenv import load_dotenv

from src.content_pipeline import Pipeline
from src.llm_integration import ContentGeneratorError

load_dotenv()

# ── Page config ───────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="SRH AI Content Creator",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Constants ─────────────────────────────────────────────────────────────────

SRH_ORANGE = "#E84E0F"

CONTENT_TYPES = {
    "blog":       ("📝", "Blog Post"),
    "social":     ("📱", "Social Media"),
    "program":    ("🎓", "Program"),
    "newsletter": ("📧", "Newsletter"),
}

TOPIC_SUGGESTIONS = {
    "blog":       ["AI Ethics at SRH", "Why Study in Berlin?", "The CORE Principle", "Careers in Data Science"],
    "social":     ["Open Day June 2026", "New AI Programme", "Student Success Story", "Berlin Campus Life"],
    "program":    ["MSc Applied Data Science and AI", "Executive MBA", "BSc Computer Science", "MSc Big Data & Analytics"],
    "newsletter": ["April Campus Updates", "New Programme Launch", "Alumni Success Stories", "Open Day Invitation"],
}

AUDIENCES = ["Prospective Students", "Current Students", "Faculty & Staff", "Industry Partners", "Alumni"]
TONES     = ["Academic", "Formal", "Professional", "Friendly", "Conversational"]
LENGTHS   = {"Short": "~300 words", "Medium": "~650 words", "Long": "~1,000 words"}

KB_DIRS = {
    "primary":   [("PRIMARY SOURCES", "knowledge_base/primary/")],
    "secondary": [("SECONDARY SOURCES", "knowledge_base/secondary/")],
    "hybrid": [
        ("PRIMARY SOURCES", "knowledge_base/primary/"),
        ("SECONDARY SOURCES", "knowledge_base/secondary/"),
    ],
}

# ── CSS ───────────────────────────────────────────────────────────────────────

# Base selector for the outer two-panel row (avoids cascading into nested columns).
_OUTER = ".block-container > [data-testid='stVerticalBlock'] > [data-testid='stHorizontalBlock']"
_LEFT  = f"{_OUTER} > [data-testid='stColumn']:first-child"
_RIGHT = f"{_OUTER} > [data-testid='stColumn']:last-child"

st.markdown(f"""
<style>
  /* ══ CHROME ══ */
  #MainMenu, footer, header {{ visibility: hidden; }}

  /* ══ VIEWPORT LOCK ══ */
  html, body {{
      height: 100vh !important;
      overflow: hidden !important;
      margin: 0;
  }}
  [data-testid="stAppViewContainer"],
  [data-testid="stMain"] {{
      height: 100vh !important;
      overflow: hidden !important;
  }}
  .block-container {{
      padding: 0 !important;
      max-width: 100% !important;
      height: 100vh !important;
      overflow: hidden !important;
  }}
  .block-container > [data-testid="stVerticalBlock"] {{
      height: 100vh !important;
      overflow: hidden !important;
      gap: 0 !important;
  }}

  /* ══ OUTER ROW ══ */
  {_OUTER} {{
      height: 100vh !important;
      gap: 0 !important;
      align-items: stretch !important;
  }}

  /* ══ LEFT PANEL ══ */
  {_LEFT} {{
      height: 100vh !important;
      overflow-y: auto !important;
      overflow-x: hidden !important;
      background: #FAFAFA !important;
      border-right: 1px solid #E8E8E8 !important;
      padding: 0.65rem 0.9rem 1rem !important;
  }}
  /* Collapse spacing between direct children of left panel */
  {_LEFT} > [data-testid="stVerticalBlock"] {{
      gap: 0.05rem !important;
  }}
  /* Remove default bottom margin on every element-container in left panel */
  {_LEFT} .element-container {{
      margin-bottom: 0 !important;
  }}

  /* ══ RIGHT PANEL ══ */
  {_RIGHT} {{
      height: 100vh !important;
      overflow-y: auto !important;
      overflow-x: hidden !important;
      background: #FFFFFF !important;
      padding: 0.65rem 1.1rem 1rem !important;
  }}

  /* ══ PANEL HEADER ══ */
  .panel-title {{
      font-size: 0.92rem;
      font-weight: 700;
      color: #1A1A1A;
      line-height: 1.2;
      margin: 0;
  }}
  .panel-sub {{
      font-size: 0.65rem;
      color: #AAA;
      margin: 0;
  }}

  /* ══ SECTION LABELS ══ */
  .mlabel {{
      font-size: 0.61rem;
      font-weight: 700;
      text-transform: uppercase;
      letter-spacing: 0.08em;
      color: #999;
      margin: 0.42rem 0 0.1rem;
      display: block;
  }}
  .mlabel-sug {{
      font-size: 0.57rem;
      font-weight: 700;
      text-transform: uppercase;
      letter-spacing: 0.08em;
      color: #BBB;
      margin: 0.28rem 0 0.1rem;
      display: block;
  }}

  /* ══ INPUTS ══ */
  [data-testid="stTextInput"] input {{
      height: 32px !important;
      font-size: 0.8rem !important;
      padding: 0 0.55rem !important;
  }}
  [data-testid="stSelectbox"] > div > div {{
      height: 32px !important;
      min-height: 32px !important;
      font-size: 0.8rem !important;
  }}

  /* ══ RADIO → PILLS ══
     Hide the dot; turn every label into a clickable pill. */
  div[data-testid="stRadio"] div[role="radiogroup"] label input[type="radio"] {{
      position: absolute !important;
      opacity: 0 !important;
      width: 0 !important;
      height: 0 !important;
      pointer-events: none !important;
  }}
  div[data-testid="stRadio"] div[role="radiogroup"] {{
      display: flex !important;
      flex-wrap: wrap !important;
      gap: 0.22rem !important;
  }}
  div[data-testid="stRadio"] div[role="radiogroup"] label {{
      display: inline-flex !important;
      align-items: center !important;
      border: 1px solid #DDD !important;
      border-radius: 4px !important;
      padding: 0 0.5rem !important;
      font-size: 0.76rem !important;
      cursor: pointer !important;
      background: #F7F7F7 !important;
      margin: 0 !important;
      transition: border-color 0.12s, background 0.12s !important;
      white-space: nowrap !important;
      height: 26px !important;
      line-height: 1 !important;
  }}
  div[data-testid="stRadio"] div[role="radiogroup"] label:has(input:checked) {{
      border-color: {SRH_ORANGE} !important;
      background: #FFF3EF !important;
      color: {SRH_ORANGE} !important;
      font-weight: 600 !important;
  }}
  div[data-testid="stRadio"] div[role="radiogroup"] label:hover {{
      border-color: #CCC !important;
      background: #F0F0F0 !important;
  }}
  div[data-testid="stRadio"] div[role="radiogroup"] label:has(input:checked):hover {{
      border-color: {SRH_ORANGE} !important;
      background: #FFF3EF !important;
  }}
  div[data-testid="stRadio"] {{ margin-bottom: 0 !important; }}

  /* ══ EXPANDER ══ */
  [data-testid="stExpander"] {{
      border: 1px solid #EEE !important;
      border-radius: 5px !important;
      margin: 0.3rem 0 0 !important;
  }}
  [data-testid="stExpander"] summary {{
      font-size: 0.75rem !important;
      font-weight: 600 !important;
      padding: 0.35rem 0.6rem !important;
      min-height: unset !important;
  }}
  [data-testid="stExpander"] > div[data-testid="stExpanderDetails"] {{
      padding: 0.3rem 0.6rem 0.5rem !important;
  }}

  /* ══ SELECT SLIDER ══ */
  [data-testid="stSlider"] {{ padding: 0.1rem 0 0.3rem !important; }}
  [data-testid="stSlider"] label {{ display: none !important; }}

  /* ══ GENERATE BUTTON ══ */
  div.stButton > button[kind="primary"] {{
      background: {SRH_ORANGE} !important;
      border: none !important;
      border-radius: 5px !important;
      font-weight: 700 !important;
      font-size: 0.85rem !important;
      height: 36px !important;
      color: white !important;
      width: 100%;
      margin-top: 0.25rem;
  }}
  div.stButton > button[kind="primary"]:hover {{
      filter: brightness(0.9) !important;
  }}

  /* ══ SUGGESTION BUTTONS ══ */
  div.stButton > button[kind="secondary"] {{
      border-radius: 4px !important;
      border: 1px solid #E0E0E0 !important;
      background: #F6F6F6 !important;
      color: #444 !important;
      font-size: 0.68rem !important;
      padding: 0 0.4rem !important;
      height: 26px !important;
      min-height: 26px !important;
      line-height: 1 !important;
  }}
  div.stButton > button[kind="secondary"]:hover {{
      border-color: {SRH_ORANGE} !important;
      background: #FFF0EC !important;
      color: {SRH_ORANGE} !important;
  }}
  [data-testid="stColumn"] div.stButton {{
      margin-bottom: 0.08rem !important;
  }}

  /* ══ TABS ══ */
  div[data-testid="stTabs"] button[role="tab"] {{
      font-weight: 600;
      font-size: 0.8rem;
      padding: 0.35rem 0.8rem;
  }}
  div[data-testid="stTabs"] button[role="tab"][aria-selected="true"] {{
      border-bottom: 2px solid {SRH_ORANGE} !important;
      color: {SRH_ORANGE} !important;
  }}

  /* ══ STATUS BADGE ══ */
  .status-badge {{
      display: inline-flex;
      align-items: center;
      gap: 0.35rem;
      background: #F0FFF4;
      border: 1px solid #86EFAC;
      border-radius: 5px;
      padding: 0.25rem 0.6rem;
      font-size: 0.72rem;
      font-weight: 600;
      color: #166534;
      margin-bottom: 0.4rem;
  }}

  /* ══ KB CARDS ══ */
  .kb-card {{
      background: #FAFAFA;
      border: 1px solid #EBEBEB;
      border-radius: 5px;
      padding: 0.3rem 0.6rem;
      margin-bottom: 0.2rem;
      display: flex;
      justify-content: space-between;
      align-items: center;
      font-size: 0.75rem;
  }}
  .tag-primary   {{ background: #DBEAFE; color: #1D4ED8; font-size: 0.6rem; font-weight: 700;
                    padding: 0.07rem 0.3rem; border-radius: 999px; margin-right: 0.35rem; }}
  .tag-secondary {{ background: #FEF3C7; color: #B45309; font-size: 0.6rem; font-weight: 700;
                    padding: 0.07rem 0.3rem; border-radius: 999px; margin-right: 0.35rem; }}

  /* ══ DIFF ══ */
  mark.unique {{ background: #BBF7D0; border-radius: 3px; padding: 0 2px; }}

  /* ══ MISC ══ */
  hr {{ margin: 0.35rem 0 !important; border-color: #EEE !important; }}
  [data-testid="stSpinner"] svg {{ stroke: {SRH_ORANGE}; }}
  [data-testid="stCaption"] {{ font-size: 0.65rem !important; color: #BBB !important; }}

  /* ══ MOBILE FALLBACK ══ */
  @media (max-width: 768px) {{
      html, body {{ overflow: auto !important; height: auto !important; }}
      {_OUTER} {{ flex-direction: column !important; height: auto !important; }}
      {_LEFT}, {_RIGHT} {{ height: auto !important; max-height: none !important; }}
  }}
</style>
""", unsafe_allow_html=True)

# ── Helpers ───────────────────────────────────────────────────────────────────

def load_baseline() -> str:
    try:
        with open(os.path.join("examples", "chatgpt_output.md"), encoding="utf-8") as f:
            raw = f.read()
        lines = [l for l in raw.splitlines() if not l.strip().startswith("<!--")]
        return "\n".join(lines).strip()
    except FileNotFoundError:
        return ""


@st.cache_resource(show_spinner=False)
def load_kb(source: str) -> tuple[str, list[dict]]:
    from src.document_processor import MarkdownProcessor
    sections, meta = [], []
    for label, kb_dir in KB_DIRS[source]:
        try:
            docs = MarkdownProcessor(kb_dir).process_all()
            block = "\n\n---\n\n".join(f"### {d['filename']}\n{d['content']}" for d in docs)
            sections.append(f"## {label}\n\n{block}")
            for d in docs:
                meta.append({
                    "filename": d["filename"],
                    "words": len(d["content"].split()),
                    "source": "primary" if "primary" in kb_dir else "secondary",
                })
        except (FileNotFoundError, ValueError):
            pass
    return "\n\n===\n\n".join(sections), meta


def make_docx(content: str, topic: str) -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading(topic, 0)
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], 3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], 2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], 1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            p = doc.add_paragraph(stripped)
            if p.runs:
                p.runs[0].font.size = Pt(11)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def highlight_unique(our: str, theirs: str) -> str:
    their_words = {w.lower() for w in re.findall(r'\b\w{5,}\b', theirs)}
    def tag(m):
        w = m.group(0)
        return f'<mark class="unique">{w}</mark>' if w.lower() not in their_words else w
    lines = []
    for line in our.splitlines():
        if line.startswith(("#", "`", "|")):
            lines.append(line)
        else:
            lines.append(re.sub(r'\b[A-Z][a-zA-Z]{4,}\b', tag, line))
    return "<br>".join(lines)


# ── Session state ─────────────────────────────────────────────────────────────

DEFAULTS = {
    "generated": None, "file_meta": [],
    "last_topic": "", "last_type": "blog",
    "last_language": "english", "last_kb_source": "hybrid",
    "last_audience": "Prospective Students",
    "last_tone": "Professional", "last_length": "Medium",
    "topic_val": "",
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ── Layout ────────────────────────────────────────────────────────────────────

left, right = st.columns([1, 1.55], gap="small")

# ════════════════════════════════════════════════════════
# LEFT PANEL — all inputs
# ════════════════════════════════════════════════════════
with left:

    # Header: logo + title side by side
    hc1, hc2 = st.columns([1, 3], gap="small")
    hc1.image("assets/srh_logo.png", width=56)
    hc2.markdown(
        '<div style="padding-top:3px">'
        '<p class="panel-title">AI Content Creator</p>'
        '<p class="panel-sub">SRH University · GPT-4o-mini</p>'
        '</div>',
        unsafe_allow_html=True,
    )

    st.markdown('<hr>', unsafe_allow_html=True)

    # ── Content type ─────────────────────────────────────
    st.markdown('<p class="mlabel">Content Type</p>', unsafe_allow_html=True)
    content_type = st.radio(
        "ctype",
        options=list(CONTENT_TYPES.keys()),
        format_func=lambda x: f"{CONTENT_TYPES[x][0]} {CONTENT_TYPES[x][1]}",
        horizontal=True,
        label_visibility="collapsed",
    )

    # ── Topic ─────────────────────────────────────────────
    st.markdown('<p class="mlabel">Topic</p>', unsafe_allow_html=True)
    topic = st.text_input(
        "topic",
        value=st.session_state.topic_val,
        placeholder=f"e.g. {TOPIC_SUGGESTIONS[content_type][0]}",
        label_visibility="collapsed",
    )

    st.markdown('<p class="mlabel-sug">Quick suggestions</p>', unsafe_allow_html=True)
    sug_cols = st.columns(2, gap="small")
    for i, sug in enumerate(TOPIC_SUGGESTIONS[content_type][:4]):
        if sug_cols[i % 2].button(sug, key=f"sug_{sug}", use_container_width=True):
            st.session_state.topic_val = sug
            st.rerun()

    st.markdown('<hr>', unsafe_allow_html=True)

    # ── Audience + Language in one row ────────────────────
    ac, lc = st.columns([3, 2], gap="small")
    with ac:
        st.markdown('<p class="mlabel">Audience</p>', unsafe_allow_html=True)
        audience = st.selectbox("audience", AUDIENCES, label_visibility="collapsed")
    with lc:
        st.markdown('<p class="mlabel">Language</p>', unsafe_allow_html=True)
        language = st.radio(
            "lang", ["english", "german"],
            format_func=lambda x: "🇬🇧 EN" if x == "english" else "🇩🇪 DE",
            horizontal=True,
            label_visibility="collapsed",
        )

    # ── Advanced options (collapsed by default) ───────────
    with st.expander("⚙️  KB · Tone · Length", expanded=False):
        st.markdown('<p class="mlabel">Knowledge Base</p>', unsafe_allow_html=True)
        kb_source = st.radio(
            "kb", ["hybrid", "primary", "secondary"],
            format_func=lambda x: {
                "hybrid": "🟢 Hybrid", "primary": "🔵 Primary", "secondary": "🟡 Secondary",
            }[x],
            horizontal=True, label_visibility="collapsed",
        )
        st.markdown('<p class="mlabel">Tone</p>', unsafe_allow_html=True)
        tone = st.select_slider(
            "tone", options=TONES, value="Professional", label_visibility="collapsed",
        )
        st.markdown('<p class="mlabel">Length</p>', unsafe_allow_html=True)
        length = st.radio(
            "length", list(LENGTHS.keys()),
            format_func=lambda x: f"{x}  ({LENGTHS[x]})",
            horizontal=True, label_visibility="collapsed",
        )

    st.markdown('<hr>', unsafe_allow_html=True)

    # ── Generate ──────────────────────────────────────────
    generate_btn = st.button(
        f"✨  Generate {CONTENT_TYPES[content_type][1]}",
        type="primary", use_container_width=True,
    )
    st.caption("Powered by GPT-4o-mini · SRH Knowledge Base")


# ════════════════════════════════════════════════════════
# RIGHT PANEL — output
# ════════════════════════════════════════════════════════
with right:

    # ── Generation trigger ────────────────────────────────
    if generate_btn:
        active_topic = st.session_state.topic_val if not topic.strip() else topic
        if not active_topic.strip():
            st.warning("Please enter a topic before generating.")
        elif not os.getenv("OPENAI_API_KEY"):
            st.error("OPENAI_API_KEY is not set. Add it to your `.env` file and restart.")
        else:
            try:
                with st.spinner("Loading SRH knowledge base…"):
                    kb_context, file_meta = load_kb(kb_source)

                style_preamble = (
                    f"GENERATION SETTINGS — follow these precisely:\n"
                    f"  Tone    : {tone}\n"
                    f"  Length  : {LENGTHS[length]}\n"
                    f"  Audience: {audience}\n\n"
                )

                with st.spinner(f"Generating {CONTENT_TYPES[content_type][1]}…"):
                    pipeline = Pipeline(language=language)
                    pipeline.kb_context = style_preamble + kb_context
                    pipeline.monitor(
                        topic=active_topic, content_type=content_type,
                        audience=audience, extra="",
                    )
                    content = pipeline.publish()

                st.session_state.generated      = content
                st.session_state.file_meta      = file_meta
                st.session_state.last_topic     = active_topic
                st.session_state.last_type      = content_type
                st.session_state.last_language  = language
                st.session_state.last_kb_source = kb_source
                st.session_state.last_audience  = audience
                st.session_state.last_tone      = tone
                st.session_state.last_length    = length

            except ContentGeneratorError as e:
                st.error(f"Generation failed:\n{e}")
            except (FileNotFoundError, ValueError) as e:
                st.error(f"Knowledge base error:\n{e}")

    # ── Output tabs ───────────────────────────────────────
    if st.session_state.generated:
        s = st.session_state
        lang_flag = "🇩🇪" if s.last_language == "german" else "🇬🇧"
        kb_dot    = {"hybrid": "🟢", "primary": "🔵", "secondary": "🟡"}[s.last_kb_source]
        st.markdown(
            f'<div class="status-badge">✅ &nbsp;{CONTENT_TYPES[s.last_type][0]} {CONTENT_TYPES[s.last_type][1]}'
            f' &nbsp;·&nbsp; {lang_flag} {s.last_language.title()}'
            f' &nbsp;·&nbsp; {kb_dot} {s.last_kb_source.title()} KB'
            f' &nbsp;·&nbsp; {s.last_tone} &nbsp;·&nbsp; {s.last_length}'
            f'</div>',
            unsafe_allow_html=True,
        )

        tab1, tab2, tab3 = st.tabs([
            "📄  Generated Content",
            "🔍  Uniqueness Analysis",
            "📚  Knowledge Base Used",
        ])

        with tab1:
            st.markdown(s.generated)
            st.divider()
            c1, c2 = st.columns(2)
            with c1:
                st.download_button(
                    "⬇️  Download .md", data=s.generated,
                    file_name=f"srh_{s.last_type}_{s.last_language}.md",
                    mime="text/markdown", use_container_width=True,
                )
            with c2:
                try:
                    docx_bytes = make_docx(s.generated, s.last_topic)
                    st.download_button(
                        "⬇️  Download .docx", data=docx_bytes,
                        file_name=f"srh_{s.last_type}_{s.last_language}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                except ImportError:
                    st.caption("Install `python-docx` for Word export.")

        with tab2:
            baseline = load_baseline()
            if not baseline:
                st.info("No ChatGPT baseline found. Add content to `examples/chatgpt_output.md`.")
            else:
                col_ours, col_theirs = st.columns(2)
                with col_ours:
                    st.markdown("**🟢 Our RAG System**")
                    st.caption("Green = terms unique to our output, absent from ChatGPT's")
                    highlighted = highlight_unique(s.generated, baseline)
                    st.markdown(highlighted, unsafe_allow_html=True)
                with col_theirs:
                    st.markdown("**⚪ ChatGPT Baseline**")
                    st.caption("Public knowledge only — no SRH context")
                    st.markdown(baseline)

                st.divider()
                st.markdown("**Factual specificity scorecard**")
                st.markdown("""
| Criterion | ChatGPT | Our System |
|---|:---:|:---:|
| CORE principle — specific details | ⚠️ Generic | ✅ |
| 5-week block / 8 ECTS structure | ❌ | ✅ |
| Named alumni quote + job title | ❌ | ✅ |
| 140+ countries statistic | ❌ | ✅ |
| Career:Skills programme named | ❌ | ✅ |
| SRH brand voice enforced | ❌ | ✅ |
| British English throughout | ❌ | ✅ |
                """)

        with tab3:
            meta = s.file_meta
            if not meta:
                st.info("No file metadata — regenerate to populate this tab.")
            else:
                primary_files   = [f for f in meta if f["source"] == "primary"]
                secondary_files = [f for f in meta if f["source"] == "secondary"]

                if primary_files:
                    st.markdown(f"**🔵 Primary Sources** — {len(primary_files)} files")
                    for f in primary_files:
                        st.markdown(
                            f'<div class="kb-card">'
                            f'<span><span class="tag-primary">PRIMARY</span>{f["filename"]}</span>'
                            f'<span style="color:#999;font-size:0.72rem">{f["words"]:,} words</span>'
                            f'</div>',
                            unsafe_allow_html=True,
                        )

                if secondary_files:
                    st.markdown(f"**🟡 Secondary Sources** — {len(secondary_files)} files")
                    for f in secondary_files:
                        st.markdown(
                            f'<div class="kb-card">'
                            f'<span><span class="tag-secondary">SECONDARY</span>{f["filename"]}</span>'
                            f'<span style="color:#999;font-size:0.72rem">{f["words"]:,} words</span>'
                            f'</div>',
                            unsafe_allow_html=True,
                        )

                total_words = sum(f["words"] for f in meta)
                st.divider()
                st.caption(f"Total context: {total_words:,} words across {len(meta)} documents")

    # ── Welcome state ─────────────────────────────────────
    else:
        st.markdown("### SRH AI Content Creator")
        st.markdown(
            "Select a content type, enter your topic, and click **Generate** — "
            "every word is grounded in SRH's own knowledge base."
        )
        st.divider()

        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown(f"""
<div style="border:1px solid #E8E8E8;border-radius:8px;padding:0.9rem;font-size:0.82rem">
<b>Why this differs from ChatGPT</b><br><br>
✅ &nbsp;Named alumni quotes (Hassan Hadeed, QESTRIT)<br>
✅ &nbsp;5-week blocks, 8 ECTS cap — exact program data<br>
✅ &nbsp;"140 countries" stat, not "diverse student body"<br>
✅ &nbsp;Career:Skills programme named correctly<br>
✅ &nbsp;CORE principle explained — not just mentioned<br>
✅ &nbsp;Brand voice enforced by prompt engineering<br>
✅ &nbsp;Bilingual — English &amp; German output<br>
</div>""", unsafe_allow_html=True)

        with col_b:
            st.markdown(f"""
<div style="border:1px solid #E8E8E8;border-radius:8px;padding:0.9rem;font-size:0.82rem">
<b>Content types</b><br><br>
📝 &nbsp;<b>Blog Post</b> — 700–1,000 words, thought leadership<br><br>
📱 &nbsp;<b>Social Media</b> — LinkedIn + Instagram, dual format<br><br>
🎓 &nbsp;<b>Program</b> — prospectus-ready description copy<br><br>
📧 &nbsp;<b>Newsletter</b> — subject line, sections &amp; CTA<br>
</div>""", unsafe_allow_html=True)
