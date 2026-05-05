"""
RAG ingestion pipeline for client/project knowledge sources.

Pipeline:
  extract text -> clean -> chunk -> embed -> store document/chunks/embeddings

The service is UI-agnostic. FastAPI passes typed inputs into this module, while
tests can use in-memory repositories and fake embedders.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from hashlib import sha256
from io import BytesIO
import logging
import os
from typing import Any, Literal, Protocol
from uuid import UUID

from dotenv import load_dotenv
from openai import OpenAI

from src.supabase_client import get_supabase_admin_client

load_dotenv()
logger = logging.getLogger(__name__)

SourceKind = Literal["brand", "product", "audience", "market", "competitor", "campaign", "policy", "other"]
ContentType = Literal["blog", "social", "email", "newsletter", "ad", "landing_page", "program", "other"]
ContentChannel = Literal["website", "linkedin", "instagram", "facebook", "x", "email", "newsletter", "ads", "blog", "other"]

SUPPORTED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/markdown",
    "text/x-markdown",
}


class RagIngestionError(ValueError):
    """Raised when a knowledge source cannot be ingested safely."""


@dataclass(frozen=True)
class KnowledgeSourceInput:
    organization_id: str
    client_id: str
    project_id: str | None
    title: str
    source_kind: SourceKind = "other"
    content_type: ContentType | None = None
    language: str | None = None
    channel: ContentChannel | None = None
    tags: list[str] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    text: str | None = None
    file_bytes: bytes | None = None
    filename: str | None = None
    mime_type: str | None = None
    uploaded_by: str | None = None


@dataclass(frozen=True)
class TextChunk:
    index: int
    content: str
    token_count: int
    chunk_hash: str


@dataclass(frozen=True)
class IngestedDocument:
    document_id: str
    duplicate: bool
    content_hash: str
    chunk_count: int
    embedding_count: int
    status: str
    message: str


class EmbeddingProvider(Protocol):
    model: str
    dimensions: int

    def embed_texts(self, texts: list[str]) -> list[list[float]]:
        """Embed a batch of chunk texts."""
        ...


class KnowledgeRepository(Protocol):
    def find_document_by_hash(self, client_id: str, project_id: str | None, content_hash: str) -> dict[str, Any] | None:
        """Return an existing document with the same tenant/project hash."""
        ...

    def create_document(self, source: KnowledgeSourceInput, text: str, content_hash: str) -> dict[str, Any]:
        """Persist document metadata."""
        ...

    def mark_document_status(self, document_id: str, status: str, error_message: str | None = None) -> None:
        """Update document ingestion status."""
        ...

    def create_chunk(self, source: KnowledgeSourceInput, document_id: str, chunk: TextChunk) -> dict[str, Any]:
        """Persist one text chunk."""
        ...

    def create_embedding(
        self,
        source: KnowledgeSourceInput,
        document_id: str,
        chunk_id: str,
        embedding: list[float],
        model: str,
        dimensions: int,
    ) -> dict[str, Any]:
        """Persist one chunk embedding."""
        ...


def validate_uuid(value: str, field_name: str) -> str:
    """Validate UUID strings before they reach Supabase."""
    try:
        UUID(value)
    except ValueError as error:
        raise RagIngestionError(f"{field_name} must be a valid UUID.") from error
    return value


def validate_source(source: KnowledgeSourceInput) -> None:
    """Validate ingestion input."""
    validate_uuid(source.organization_id, "organization_id")
    validate_uuid(source.client_id, "client_id")
    if source.project_id:
        validate_uuid(source.project_id, "project_id")
    if source.uploaded_by:
        validate_uuid(source.uploaded_by, "uploaded_by")
    if not source.title.strip():
        raise RagIngestionError("title is required.")
    if not source.text and not source.file_bytes:
        raise RagIngestionError("Provide either pasted text or an uploaded file.")
    if source.file_bytes and not source.filename:
        raise RagIngestionError("filename is required for uploaded files.")


def extract_text_from_source(source: KnowledgeSourceInput) -> str:
    """Extract raw text from pasted text or supported file bytes."""
    validate_source(source)
    if source.text is not None:
        logger.info("Extracting text from pasted source title='%s'", source.title)
        return source.text

    assert source.file_bytes is not None
    filename = source.filename or "upload"
    mime_type = source.mime_type or detect_mime_type(filename)
    logger.info("Extracting text from file filename=%s mime_type=%s bytes=%s", filename, mime_type, len(source.file_bytes))

    if mime_type == "application/pdf" or filename.lower().endswith(".pdf"):
        return extract_pdf_text(source.file_bytes)
    if (
        mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or filename.lower().endswith(".docx")
    ):
        return extract_docx_text(source.file_bytes)
    if mime_type in {"text/plain", "text/markdown", "text/x-markdown"} or filename.lower().endswith((".txt", ".md", ".markdown")):
        return source.file_bytes.decode("utf-8")

    raise RagIngestionError(
        f"Unsupported file type: {mime_type}. Supported sources: PDF, DOCX, TXT, Markdown, and pasted text."
    )


def detect_mime_type(filename: str) -> str:
    """Infer a conservative MIME type from filename extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return "application/pdf"
    if lower.endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if lower.endswith((".md", ".markdown")):
        return "text/markdown"
    if lower.endswith(".txt"):
        return "text/plain"
    return "application/octet-stream"


def extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from a PDF using pypdf."""
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as error:
        raise RagIngestionError("PDF ingestion requires pypdf. Install dependencies with pip install -r requirements.txt.") from error

    reader = PdfReader(BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(page.strip() for page in pages if page.strip())


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract paragraph text from a DOCX file."""
    try:
        from docx import Document
    except ModuleNotFoundError as error:
        raise RagIngestionError("DOCX ingestion requires python-docx. Install dependencies with pip install -r requirements.txt.") from error

    document = Document(BytesIO(file_bytes))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n\n".join(paragraphs)


def clean_text(text: str) -> str:
    """Normalize whitespace while preserving paragraph boundaries."""
    normalized = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [" ".join(line.strip().split()) for line in normalized.split("\n")]

    paragraphs: list[str] = []
    current: list[str] = []
    for line in lines:
        if line:
            current.append(line)
            continue
        if current:
            paragraphs.append(" ".join(current))
            current = []
    if current:
        paragraphs.append(" ".join(current))

    return "\n\n".join(paragraphs).strip()


def hash_text(text: str) -> str:
    """Return a stable SHA-256 hash for duplicate detection."""
    return sha256(text.encode("utf-8")).hexdigest()


def count_tokens_roughly(text: str) -> int:
    """Approximate token count without adding tokenizer dependencies."""
    return max(1, int(len(text.split()) * 1.33))


def chunk_text(text: str, max_tokens: int = 420, overlap_tokens: int = 60) -> list[TextChunk]:
    """Chunk text by paragraphs, then by words when paragraphs are too large."""
    if max_tokens <= overlap_tokens:
        raise RagIngestionError("max_tokens must be greater than overlap_tokens.")

    paragraphs = [paragraph.strip() for paragraph in text.split("\n\n") if paragraph.strip()]
    if not paragraphs:
        raise RagIngestionError("No usable text found after cleaning.")

    chunks: list[str] = []
    current: list[str] = []
    current_tokens = 0

    for paragraph in paragraphs:
        paragraph_tokens = count_tokens_roughly(paragraph)
        if paragraph_tokens > max_tokens:
            if current:
                chunks.append("\n\n".join(current))
                current = []
                current_tokens = 0
            chunks.extend(split_large_paragraph(paragraph, max_tokens, overlap_tokens))
            continue

        if current and current_tokens + paragraph_tokens > max_tokens:
            chunks.append("\n\n".join(current))
            overlap = trailing_words(" ".join(current), overlap_tokens)
            current = [overlap, paragraph] if overlap else [paragraph]
            current_tokens = count_tokens_roughly(" ".join(current))
        else:
            current.append(paragraph)
            current_tokens += paragraph_tokens

    if current:
        chunks.append("\n\n".join(current))

    return [
        TextChunk(index=index, content=content, token_count=count_tokens_roughly(content), chunk_hash=hash_text(content))
        for index, content in enumerate(chunks)
        if content.strip()
    ]


def split_large_paragraph(paragraph: str, max_tokens: int, overlap_tokens: int) -> list[str]:
    """Split an oversized paragraph with word overlap."""
    words = paragraph.split()
    approx_words = max(80, int(max_tokens / 1.33))
    overlap_words = max(0, int(overlap_tokens / 1.33))
    step = max(1, approx_words - overlap_words)
    return [" ".join(words[start : start + approx_words]) for start in range(0, len(words), step)]


def trailing_words(text: str, approx_tokens: int) -> str:
    """Return trailing words approximating the desired overlap size."""
    word_count = max(0, int(approx_tokens / 1.33))
    if word_count == 0:
        return ""
    return " ".join(text.split()[-word_count:])


class OpenAIEmbeddingProvider:
    """OpenAI embedding provider for RAG chunks."""

    def __init__(self, api_key: str | None = None, model: str | None = None, dimensions: int = 1536):
        api_key = api_key or os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise RagIngestionError("OPENAI_API_KEY is required to generate embeddings.")
        self.client = OpenAI(api_key=api_key)
        self.model = model or os.getenv("EMBEDDING_MODEL", "text-embedding-3-small")
        self.dimensions = dimensions

    def embed_texts(self, texts: list[str]) -> list[list[float]]:
        """Generate embeddings for a batch of chunk texts."""
        if not texts:
            return []
        logger.info("Generating embeddings chunks=%s model=%s", len(texts), self.model)
        response = self.client.embeddings.create(model=self.model, input=texts)
        embeddings = [item.embedding for item in response.data]
        if len(embeddings) != len(texts):
            raise RagIngestionError("Embedding provider returned a different number of embeddings than inputs.")
        return embeddings


class SupabaseKnowledgeRepository:
    """Supabase persistence adapter for RAG documents, chunks, and embeddings."""

    def __init__(self, client: Any | None = None):
        self.client = client or get_supabase_admin_client()

    def find_document_by_hash(self, client_id: str, project_id: str | None, content_hash: str) -> dict[str, Any] | None:
        query = (
            self.client.table("uploaded_documents")
            .select("id,status,title,content_hash")
            .eq("client_id", client_id)
            .eq("content_hash", content_hash)
            .limit(1)
        )
        if project_id:
            query = query.eq("project_id", project_id)
        else:
            query = query.is_("project_id", "null")
        response = query.execute()
        return response.data[0] if response.data else None

    def create_document(self, source: KnowledgeSourceInput, text: str, content_hash: str) -> dict[str, Any]:
        payload = {
            "organization_id": source.organization_id,
            "client_id": source.client_id,
            "project_id": source.project_id,
            "uploaded_by": source.uploaded_by,
            "title": source.title,
            "original_filename": source.filename or source.title,
            "mime_type": source.mime_type,
            "file_size_bytes": len(source.file_bytes) if source.file_bytes else len(text.encode("utf-8")),
            "content_hash": content_hash,
            "source_kind": source.source_kind,
            "status": "processing",
            "language": source.language,
            "content_type": source.content_type,
            "channel": source.channel,
            "tags": source.tags,
            "metadata": {**source.metadata, "ingestion_source": "file" if source.file_bytes else "pasted_text"},
        }
        response = self.client.table("uploaded_documents").insert(payload).execute()
        return response.data[0]

    def mark_document_status(self, document_id: str, status: str, error_message: str | None = None) -> None:
        payload = {"status": status, "error_message": error_message}
        self.client.table("uploaded_documents").update(payload).eq("id", document_id).execute()

    def create_chunk(self, source: KnowledgeSourceInput, document_id: str, chunk: TextChunk) -> dict[str, Any]:
        payload = {
            "organization_id": source.organization_id,
            "client_id": source.client_id,
            "project_id": source.project_id,
            "document_id": document_id,
            "chunk_index": chunk.index,
            "content": chunk.content,
            "chunk_hash": chunk.chunk_hash,
            "token_count": chunk.token_count,
            "source_kind": source.source_kind,
            "language": source.language,
            "content_type": source.content_type,
            "channel": source.channel,
            "tags": source.tags,
            "metadata": source.metadata,
        }
        response = self.client.table("document_chunks").insert(payload).execute()
        return response.data[0]

    def create_embedding(
        self,
        source: KnowledgeSourceInput,
        document_id: str,
        chunk_id: str,
        embedding: list[float],
        model: str,
        dimensions: int,
    ) -> dict[str, Any]:
        payload = {
            "organization_id": source.organization_id,
            "client_id": source.client_id,
            "project_id": source.project_id,
            "document_id": document_id,
            "chunk_id": chunk_id,
            "embedding_model": model,
            "embedding_dimensions": dimensions,
            "embedding": embedding,
        }
        response = self.client.table("document_embeddings").insert(payload).execute()
        return response.data[0]

    def match_chunks(
        self,
        query_embedding: list[float],
        client_id: str | None = None,
        project_id: str | None = None,
        content_type: str | None = None,
        language: str | None = None,
        channel: str | None = None,
        match_count: int = 8,
        match_threshold: float = 0.72,
    ) -> list[dict[str, Any]]:
        """Call the Supabase pgvector similarity-search RPC."""
        response = self.client.rpc(
            "match_document_chunks",
            {
                "query_embedding": query_embedding,
                "match_count": match_count,
                "match_threshold": match_threshold,
                "filter_client_id": client_id,
                "filter_project_id": project_id,
                "filter_content_type": content_type,
                "filter_language": language,
                "filter_channel": channel,
            },
        ).execute()
        return list(response.data or [])


class RagIngestionService:
    """Coordinates extraction, cleaning, chunking, embedding, and storage."""

    def __init__(self, repository: KnowledgeRepository | None = None, embedder: EmbeddingProvider | None = None):
        self.repository = repository or SupabaseKnowledgeRepository()
        self.embedder = embedder or OpenAIEmbeddingProvider()

    def ingest(self, source: KnowledgeSourceInput) -> IngestedDocument:
        logger.info("Starting RAG ingestion title='%s' client_id=%s project_id=%s", source.title, source.client_id, source.project_id)
        raw_text = extract_text_from_source(source)
        cleaned = clean_text(raw_text)
        if len(cleaned) < 20:
            raise RagIngestionError("Source text is too short after cleaning. Add more useful source content.")

        content_hash = hash_text(cleaned)
        duplicate = self.repository.find_document_by_hash(source.client_id, source.project_id, content_hash)
        if duplicate:
            logger.info("Duplicate knowledge source detected document_id=%s hash=%s", duplicate["id"], content_hash)
            return IngestedDocument(
                document_id=str(duplicate["id"]),
                duplicate=True,
                content_hash=content_hash,
                chunk_count=0,
                embedding_count=0,
                status=str(duplicate.get("status", "ready")),
                message="Duplicate source detected. Existing document was reused.",
            )

        document = self.repository.create_document(source, cleaned, content_hash)
        document_id = str(document["id"])
        try:
            chunks = chunk_text(cleaned)
            embeddings = self.embedder.embed_texts([chunk.content for chunk in chunks])
            for chunk, embedding in zip(chunks, embeddings):
                chunk_record = self.repository.create_chunk(source, document_id, chunk)
                self.repository.create_embedding(
                    source,
                    document_id,
                    str(chunk_record["id"]),
                    embedding,
                    self.embedder.model,
                    self.embedder.dimensions,
                )

            self.repository.mark_document_status(document_id, "ready")
            logger.info("Completed RAG ingestion document_id=%s chunks=%s", document_id, len(chunks))
            return IngestedDocument(
                document_id=document_id,
                duplicate=False,
                content_hash=content_hash,
                chunk_count=len(chunks),
                embedding_count=len(embeddings),
                status="ready",
                message="Knowledge source ingested successfully.",
            )
        except Exception as error:
            self.repository.mark_document_status(document_id, "failed", str(error))
            logger.exception("RAG ingestion failed document_id=%s", document_id)
            raise


def search_knowledge_chunks(
    query: str,
    client_id: str | None = None,
    project_id: str | None = None,
    content_type: str | None = None,
    language: str | None = None,
    channel: str | None = None,
    match_count: int = 8,
    match_threshold: float = 0.72,
    repository: SupabaseKnowledgeRepository | None = None,
    embedder: EmbeddingProvider | None = None,
) -> list[dict[str, Any]]:
    """Embed a retrieval query and return matching chunks from Supabase."""
    if not query.strip():
        raise RagIngestionError("query is required.")
    if client_id:
        validate_uuid(client_id, "client_id")
    if project_id:
        validate_uuid(project_id, "project_id")

    active_embedder = embedder or OpenAIEmbeddingProvider()
    active_repository = repository or SupabaseKnowledgeRepository()
    embedding = active_embedder.embed_texts([query])[0]
    return active_repository.match_chunks(
        query_embedding=embedding,
        client_id=client_id,
        project_id=project_id,
        content_type=content_type,
        language=language,
        channel=channel,
        match_count=match_count,
        match_threshold=match_threshold,
    )
