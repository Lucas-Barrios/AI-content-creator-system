"""
Backend utilities for generated-content artifacts.

These functions used to sit behind UI callbacks. They are now framework-neutral
so API routes, CLIs, or future workers can reuse them without depending on a
frontend library.
"""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
import logging
import re

from src.io_helpers import read_text_without_metadata

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class UniquenessComparison:
    baseline: str
    unique_terms: list[str]
    unique_term_count: int
    compared_word_count: int


def load_baseline(path: str = "examples/chatgpt_output.md") -> str:
    """Load the public ChatGPT baseline used for comparison."""
    logger.info("Loading comparison baseline from %s", path)
    try:
        baseline = read_text_without_metadata(path)
        logger.info("Loaded comparison baseline from %s chars=%s", path, len(baseline))
        return baseline
    except FileNotFoundError:
        logger.warning("Comparison baseline not found at %s", path)
        return ""


def compare_uniqueness(content: str, baseline: str | None = None) -> UniquenessComparison:
    """Compare generated content against a baseline and return unique terms."""
    logger.info("Comparing content uniqueness content_chars=%s", len(content))
    baseline_text = baseline if baseline is not None else load_baseline()
    baseline_words = {word.lower() for word in re.findall(r"\b\w{5,}\b", baseline_text)}
    content_words = re.findall(r"\b[A-Z][a-zA-Z]{4,}\b", content)
    unique_terms = sorted({word for word in content_words if word.lower() not in baseline_words})

    comparison = UniquenessComparison(
        baseline=baseline_text,
        unique_terms=unique_terms,
        unique_term_count=len(unique_terms),
        compared_word_count=len(content.split()),
    )
    logger.info("Completed uniqueness comparison unique_terms=%s compared_words=%s", comparison.unique_term_count, comparison.compared_word_count)
    return comparison


def make_docx_bytes(content: str, topic: str) -> bytes:
    """Create a Word document from generated Markdown-like content."""
    logger.info("Creating DOCX topic='%s' content_chars=%s", topic, len(content))
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(topic or "SRH Content", 0)

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], 3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], 2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], 1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            paragraph = doc.add_paragraph(stripped)
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(11)

    buffer = BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    logger.info("Created DOCX topic='%s' bytes=%s", topic, len(docx_bytes))
    return docx_bytes
